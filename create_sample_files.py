"""
Sample File Creator
This script creates sample Excel file and Word template for testing.
Run this file first to generate the required input files.
"""

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_sample_excel():
    """Create sample students.xlsx file with test data"""
    
    # Sample student data
    data = {
        'Name': [
            'John Smith',
            'Emily Johnson',
            'Michael Williams',
            'Sarah Brown',
            'David Jones'
        ],
        'Email': [
            'john.smith@email.com',
            'emily.johnson@email.com',
            'michael.williams@email.com',
            'sarah.brown@email.com',
            'david.jones@email.com'
        ],
        'Domain': [
            'Web Development',
            'Data Science',
            'Machine Learning',
            'Android Development',
            'Cloud Computing'
        ],
        'Duration': [
            '6 months',
            '3 months',
            '6 months',
            '4 months',
            '5 months'
        ],
        'Start Date': [
            '2024-01-15',
            '2024-02-01',
            '2024-01-20',
            '2024-03-01',
            '2024-02-15'
        ]
    }
    
    # Create DataFrame
    df = pd.DataFrame(data)
    
    # Save to Excel
    df.to_excel('students.xlsx', index=False, engine='openpyxl')
    print("✓ Created students.xlsx successfully!")
    print(f"  - Contains {len(df)} student records")
    
    return df

def create_word_template():
    """Create sample offer_template.docx with placeholders"""
    
    # Create a new Document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add Company Header
    header = doc.add_heading('TECH CORP INDUSTRIES', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Company address
    address = doc.add_paragraph('123 Tech Street, Silicon Valley, CA 94000')
    address.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add a line
    doc.add_paragraph('_' * 60)
    
    # Date placeholder
    date_para = doc.add_paragraph()
    date_para.add_run('Date: ').bold = True
    date_para.add_run('________________')
    
    doc.add_paragraph()  # Empty line
    
    # Offer Letter Title
    title = doc.add_heading('INTERNSHIP OFFER LETTER', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Empty line
    
    # Salutation
    greeting = doc.add_paragraph()
    greeting.add_run('Dear ').bold = True
    greeting.add_run('{{name}}')
    greeting.add_run(',')
    
    doc.add_paragraph()  # Empty line
    
    # Main content
    content = doc.add_paragraph()
    content.add_run('We are pleased to offer you an internship position at Tech Corp Industries. ')
    content.add_run('We were impressed with your background and believe you will be a valuable ')
    content.add_run('addition to our team.')
    
    doc.add_paragraph()  # Empty line
    
    # Internship Details
    details = doc.add_paragraph()
    details.add_run('Internship Details:').bold = True
    doc.add_paragraph()
    
    # Domain
    domain_para = doc.add_paragraph()
    domain_para.add_run('  • Domain: ').bold = True
    domain_para.add_run('{{domain}}')
    
    # Duration
    duration_para = doc.add_paragraph()
    duration_para.add_run('  • Duration: ').bold = True
    duration_para.add_run('{{duration}}')
    
    # Start Date
    start_para = doc.add_paragraph()
    start_para.add_run('  • Start Date: ').bold = True
    start_para.add_run('{{start_date}}')
    
    doc.add_paragraph()  # Empty line
    
    # Additional content
    additional = doc.add_paragraph()
    additional.add_run('This internship is a great opportunity to gain practical experience ')
    additional.add_run('in your chosen field. You will be working with our experienced team ')
    additional.add_run('and participating in real-world projects.')
    
    doc.add_paragraph()  # Empty line
    
    # Closing
    closing = doc.add_paragraph()
    closing.add_run('We look forward to welcoming you to our team!')
    
    doc.add_paragraph()  # Empty line
    doc.add_paragraph()  # Empty line
    
    # Signature
    signature = doc.add_paragraph()
    signature.add_run('Sincerely,').bold = True
    
    doc.add_paragraph()  # Empty line
    doc.add_paragraph()  # Empty line
    
    # Signature name
    sign_name = doc.add_paragraph()
    sign_name.add_run('______________________').bold = True
    sign_name.add_run('\nHR Manager\nTech Corp Industries')
    
    # Save the document
    doc.save('offer_template.docx')
    print("✓ Created offer_template.docx successfully!")
    print("  - Contains placeholders: {{name}}, {{domain}}, {{duration}}, {{start_date}}")

def main():
    """Main function to create all sample files"""
    print("Creating sample files for Internship Offer Letter Automation...\n")
    
    # Create sample Excel file
    create_sample_excel()
    print()
    
    # Create Word template
    create_word_template()
    print()
    
    print("=" * 50)
    print("All sample files created successfully!")
    print("You can now run: python gui_app.py")
    print("=" * 50)

if __name__ == '__main__':
    main()

